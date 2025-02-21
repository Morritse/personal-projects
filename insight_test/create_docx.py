from docx import Document
from docx.shared import Inches

def create_sample_docx():
    doc = Document()
    
    # Add title
    doc.add_heading('Financial Analysis Report', 0)
    
    # Add summary
    doc.add_paragraph('Executive Summary')
    summary = doc.add_paragraph(
        'This report analyzes the company\'s financial performance for Q4 2023. '
        'Overall, the company shows strong growth with improving metrics across key areas.'
    )
    
    # Add key metrics section
    doc.add_heading('Key Financial Metrics', level=1)
    metrics = [
        ('Revenue', '$1,500,000'),
        ('Net Income', '$300,000'),
        ('Profit Margin', '20%'),
        ('ROI', '15%')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    # Add risk analysis
    doc.add_heading('Risk Analysis', level=1)
    risks = doc.add_paragraph()
    risks.add_run('High Risk Areas:\n').bold = True
    doc.add_paragraph('• Market volatility\n• Supply chain disruptions\n• Regulatory changes')
    
    # Add recommendations
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph(
        '1. Increase cash reserves to improve liquidity\n'
        '2. Diversify supplier base to reduce supply chain risk\n'
        '3. Implement cost reduction measures\n'
        '4. Expand into new markets'
    )
    
    # Save the document
    doc.save('test_files/financial_report.docx')

if __name__ == '__main__':
    create_sample_docx()
